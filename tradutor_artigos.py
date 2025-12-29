"""
tradutor_artigos.py
Script simples para traduzir um arquivo .docx usando o Azure Translator Text API.

Uso:
  - Configure as variáveis de ambiente: AZURE_TRANSLATOR_KEY, AZURE_TRANSLATOR_ENDPOINT e AZURE_TRANSLATOR_REGION
  - Instale dependências: pip install -r requirements.txt
  - Execute: python tradutor_artigos.py -i "caminho/para/arquivo.docx" --to pt-br

Não coloque sua chave diretamente no código. Use variáveis de ambiente.
"""

import argparse
import logging
import os
import sys
import uuid
from typing import Optional

import requests
from docx import Document


logger = logging.getLogger(__name__)


def translate_text(text: str, target_language: str, endpoint: str, subscription_key: str, region: Optional[str] = None) -> str:
    """Traduz um texto usando o Azure Translator Text API.

    Args:
        text: Texto em inglês a ser traduzido.
        target_language: Código do idioma de destino (ex: 'pt-br').
        endpoint: URL do endpoint do serviço (ex: 'https://api.cognitive.microsofttranslator.com').
        subscription_key: Chave de subscrição (defina em AZURE_TRANSLATOR_KEY).
        region: Região (defina em AZURE_TRANSLATOR_REGION, opcional dependendo da sua subscrição).

    Returns:
        Texto traduzido.
    """
    if not text:
        return ""

    path = "/translate"
    constructed_url = endpoint.rstrip("/") + path

    headers = {
        'Ocp-Apim-Subscription-Key': subscription_key,
        'Content-type': 'application/json',
        'X-ClientTraceId': str(uuid.uuid4())
    }
    if region:
        headers['Ocp-Apim-Subscription-Region'] = region

    params = {
        'api-version': '3.0',
        'from': 'en',
        'to': target_language
    }

    body = [{'text': text}]

    try:
        resp = requests.post(constructed_url, params=params, headers=headers, json=body, timeout=15)
        resp.raise_for_status()
        data = resp.json()
        # Estrutura: [ { 'translations': [ { 'text': '...', 'to': 'pt' } ] } ]
        return data[0]['translations'][0]['text']
    except requests.RequestException as e:
        logger.error("Erro na requisição de tradução: %s", e)
        raise
    except (KeyError, IndexError) as e:
        logger.error("Resposta inesperada da API: %s", e)
        raise


def translate_document(input_path: str, target_language: str, endpoint: str, subscription_key: str, region: Optional[str] = None, output_path: Optional[str] = None) -> str:
    """Traduz um arquivo .docx inteiro (parágrafo por parágrafo) e salva uma cópia traduzida.

    Args:
        input_path: Caminho para o arquivo .docx de entrada.
        target_language: Código do idioma de destino (ex: 'pt-br').
        endpoint: URL do endpoint do serviço.
        subscription_key: Chave de subscrição.
        region: Região (opcional).
        output_path: Caminho de saída opcional. Se não fornecido, será criado a partir do nome original.

    Returns:
        Caminho do arquivo traduzido salvo.
    """
    if not os.path.isfile(input_path):
        raise FileNotFoundError(f"Arquivo não encontrado: {input_path}")

    doc = Document(input_path)

    translated_doc = Document()

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            # Mantém quebras/linhas vazias
            translated_doc.add_paragraph("")
            continue
        translated = translate_text(text, target_language, endpoint, subscription_key, region)
        translated_doc.add_paragraph(translated)

    if output_path:
        out_path = output_path
    else:
        base, ext = os.path.splitext(input_path)
        out_path = f"{base}_{target_language}{ext}"

    translated_doc.save(out_path)
    return out_path


def main(argv):
    parser = argparse.ArgumentParser(description="Traduz .docx usando Azure Translator Text API")
    parser.add_argument('-i', '--input', required=True, help='Arquivo .docx de entrada')
    parser.add_argument('-t', '--to', default='pt-br', help='Idioma destino (padrão: pt-br)')
    parser.add_argument('--endpoint', default=os.environ.get('AZURE_TRANSLATOR_ENDPOINT'), help='Endpoint do Azure (ou use AZURE_TRANSLATOR_ENDPOINT)')
    parser.add_argument('--key', default=os.environ.get('AZURE_TRANSLATOR_KEY'), help='Chave de subscrição (ou use AZURE_TRANSLATOR_KEY)')
    parser.add_argument('--region', default=os.environ.get('AZURE_TRANSLATOR_REGION'), help='Região (se aplicável)')
    parser.add_argument('-o', '--output', help='Caminho do arquivo de saída (opcional)')
    parser.add_argument('--verbose', action='store_true', help='Ativa logs detalhados')

    args = parser.parse_args(argv)

    logging.basicConfig(level=logging.DEBUG if args.verbose else logging.INFO, format='%(levelname)s: %(message)s')

    if not args.key:
        logger.error('Chave não encontrada. Defina a variável de ambiente AZURE_TRANSLATOR_KEY ou passe com --key')
        sys.exit(1)
    if not args.endpoint:
        logger.error('Endpoint não encontrado. Defina a variável de ambiente AZURE_TRANSLATOR_ENDPOINT ou passe com --endpoint')
        sys.exit(1)

    try:
        out = translate_document(args.input, args.to, args.endpoint, args.key, args.region, args.output)
        logger.info('Arquivo traduzido salvo em: %s', out)
    except Exception as e:
        logger.error('Falha ao traduzir: %s', e)
        sys.exit(1)


if __name__ == '__main__':
    main(sys.argv[1:])
